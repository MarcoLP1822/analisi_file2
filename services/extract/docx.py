"""
Estrattore DOCX
===============
Contiene:
• extract_docx_properties
• extract_docx_detailed_analysis
Usa python-docx.
"""

from __future__ import annotations

import io
from typing import Any

from docx import Document  # funzione factory
from docx.document import Document as _DocxDocument  # vero type per mypy
from docx.shared import Length

from models import DetailedDocumentAnalysis, FontInfo, ImageInfo


# ------------------------------------------------------------------ #
# helpers tipo-safe
# ------------------------------------------------------------------ #
def _emu_to_cm(val: Length | None) -> float:
    """Converte da EMU a centimetri (1 cm = 360 000 EMU)."""
    if val is None:
        return 0.0
    # python-docx Length supporta divisione per int → float
    return float(val) / 360000


def _len_or_default(val: Length | None) -> float:  # alias leggibile
    return _emu_to_cm(val)


# ------------------------------------------------------------------ #
def extract_docx_properties(file_content: bytes) -> dict[str, Any]:
    """
    Estrae:
    • dimensione pagina + margini
    • intestazioni / piè di pagina
    • TOC (heading×livello)
    • inconsistenze formato sezione-sezione
    """
    # NB: Document() restituisce _DocxDocument (vero type)
    doc: _DocxDocument = Document(io.BytesIO(file_content))  # type: ignore[call-arg]

    # ---------- dimensioni + sezioni -------------------------------
    section_data: list[dict[str, Any]] = []
    inconsistent_sections: list[dict[str, str]] = []

    for i, section in enumerate(doc.sections):
        width_cm = _len_or_default(section.page_width)
        height_cm = _len_or_default(section.page_height)
        section_info = {
            "section_num": i + 1,
            "width_cm": width_cm,
            "height_cm": height_cm,
            "margins": {
                "top_cm": _len_or_default(section.top_margin),
                "bottom_cm": _len_or_default(section.bottom_margin),
                "left_cm": _len_or_default(section.left_margin),
                "right_cm": _len_or_default(section.right_margin),
            },
        }
        section_data.append(section_info)

    first = section_data[0]
    ref_w, ref_h = first["width_cm"], first["height_cm"]
    margins = first["margins"]

    for info in section_data[1:]:
        if abs(info["width_cm"] - ref_w) > 0.1 or abs(info["height_cm"] - ref_h) > 0.1:
            inconsistent_sections.append(
                {
                    "section": info["section_num"],
                    "size": f"{info['width_cm']:.1f}x{info['height_cm']:.1f}cm",
                }
            )

    # ---------- header / footnote ----------------------------------
    headers: list[str] = []
    footnotes_texts: list[str] = []

    for section in doc.sections:
        if section.header:  # può essere None
            text = "\n".join(p.text for p in section.header.paragraphs if p.text.strip())
            if text:
                headers.append(text)

    fp = getattr(doc.part, "footnotes_part", None)
    if fp:
        for fn in fp.footnotes:
            txt = "\n".join(p.text for p in fn.paragraphs if p.text.strip())
            if txt:
                footnotes_texts.append(txt)

    # ---------- TOC / headings -------------------------------------
    headings = [p.text for p in doc.paragraphs if p.style and p.style.name.startswith("Heading")]
    has_toc = bool(headings)

    detailed_analysis = extract_docx_detailed_analysis(doc)

    return {
        "page_size": {"width_cm": ref_w, "height_cm": ref_h},
        "margins": margins,
        "has_toc": has_toc,
        "headings": headings,
        "headers": headers,
        "footnotes": footnotes_texts,
        "detailed_analysis": detailed_analysis,
        "all_section_data": section_data,
        "inconsistent_sections": inconsistent_sections,
        "has_size_inconsistencies": bool(inconsistent_sections),
    }


# ------------------------------------------------------------------ #
def extract_docx_detailed_analysis(doc: _DocxDocument) -> DetailedDocumentAnalysis:
    """
    Font + immagini + colore ecc. (codice in gran parte invariato,
    con qualche assert per soddisfare mypy).
    """
    fonts: dict[str, FontInfo] = {}
    paragraph_count = 0
    line_spacing: dict[str, float] = {}
    toc_structure: list[dict[str, str]] = []
    has_color_text = False
    colored_elements_count = 0

    # potrebbero mancare → usa default
    normal_style = doc.styles.get("Normal", None)
    default_font_name = getattr(normal_style, "font", None)
    default_font_name = default_font_name.name if default_font_name else "Default"
    if normal_style and normal_style.font and normal_style.font.size is not None:
        default_font_size: float = normal_style.font.size.pt  # type: ignore[assignment]
    else:
        default_font_size = 11.0


    for par in doc.paragraphs:
        paragraph_count += 1

        if par._element.pPr is not None and par._element.pPr.spacing is not None:
            if par._element.pPr.spacing.line is not None:
                style = par.style.name if par.style else "Unknown"
                spacing = par._element.pPr.spacing.line / 240
                line_spacing[style] = (
                    spacing if style not in line_spacing else (line_spacing[style] + spacing) / 2
                )

        if par.style and par.style.name.startswith("Heading"):
            lev = int(par.style.name.replace("Heading", "")) if par.style.name != "Heading" else 1
            toc_structure.append({"level": str(lev), "text": par.text})

        for run in par.runs:
            fname = (
                run.font.name
                or (par.style.font.name if par.style else None)
                or default_font_name
            )
            if run.font.size:
                fsize = round(run.font.size.pt, 1)
            elif par.style and par.style.font.size:
                fsize = round(par.style.font.size.pt, 1)
            else:
                fsize = round(default_font_size, 1)

            if run.font.color and run.font.color.rgb and run.font.color.rgb != "000000":
                has_color_text = True
                colored_elements_count += 1

            fi = fonts.setdefault(fname, FontInfo(sizes=[], count=0, size_counts={}))
            fi.count += 1
            if fsize not in fi.sizes:
                fi.sizes.append(fsize)
            fi.size_counts[fsize] = fi.size_counts.get(fsize, 0) + 1

    # immagini
    img_cnt = 0
    img_size = 0
    has_color_pages = False
    for rel in doc.part.rels.values():
        if rel.target_ref.startswith("media/"):
            img_cnt += 1
            if hasattr(rel.target_part, "blob"):
                img_size += len(rel.target_part.blob)
            has_color_pages = True
            colored_elements_count += 1

    img_info = (
        ImageInfo(count=img_cnt, avg_size_kb=round((img_size / img_cnt) / 1024, 2))
        if img_cnt
        else None
    )

    # metadati
    md: dict[str, str] = {}
    cp = doc.core_properties
    if cp.author:
        md["author"] = cp.author
    if cp.title:
        md["title"] = cp.title
    if cp.created:
        md["created"] = cp.created.isoformat()
    if cp.modified:
        md["modified"] = cp.modified.isoformat()

    return DetailedDocumentAnalysis(
        fonts=fonts,
        images=img_info,
        line_spacing=line_spacing,
        paragraph_count=paragraph_count,
        toc_structure=toc_structure,
        metadata=md,
        has_color_pages=has_color_pages,
        has_color_text=has_color_text,
        colored_elements_count=colored_elements_count,
    )
